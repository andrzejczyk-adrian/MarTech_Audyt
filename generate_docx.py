"""
generate_docx.py — Generyczny konwerter audytu GA4 (MD → DOCX)

Użycie:
    python generate_docx.py <ścieżka_do_audytu.md>

Jeśli brak argumentu — szuka jedynego pliku *_audyt.md w bieżącym katalogu.
Wynik zapisuje jako <nazwa_audytu>.docx obok pliku wejściowego.
"""

import sys
import re
import os
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────
# HELPERS — styl i formatowanie
# ──────────────────────────────────────────────

def _shd(cell, fill_hex):
    """Ustawia kolor tła komórki tabeli."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def status_color(text):
    t = str(text)
    if any(x in t for x in ['✅', 'OK', 'Tak', 'Spełniony']):
        return RGBColor(0x10, 0x7C, 0x10)   # zielony
    if any(x in t for x in ['❌', 'Błąd', 'Krytyczny', 'NIEMOŻLIWE', 'Nie']):
        return RGBColor(0xC0, 0x00, 0x00)   # czerwony
    if any(x in t for x in ['⚠️', 'Wymaga', 'Częściowo', 'Borderline', 'Uwaga']):
        return RGBColor(0xBF, 0x8F, 0x00)   # żółty
    if any(x in t for x in ['🔒', '➖', 'N/A', 'Brak danych']):
        return RGBColor(0x70, 0x70, 0x70)   # szary
    return None


def _add_mixed_run(para, text, base_size=10.5):
    """Dodaje tekst z obsługą **bold**, `code` i _italic_ w ramach paragrafu."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|_[^_]+_)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(base_size)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(base_size)
        else:
            if part:
                run = para.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(base_size)
                col = status_color(part)
                if col:
                    run.font.color.rgb = col
    return para


# ──────────────────────────────────────────────
# BUILDER — metody dodające elementy do doc
# ──────────────────────────────────────────────

class DocBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10.5)
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def heading(self, text, level):
        COLORS = {
            1: RGBColor(0x1F, 0x49, 0x7D),
            2: RGBColor(0x2E, 0x74, 0xB5),
            3: RGBColor(0x1F, 0x49, 0x7D),
            4: RGBColor(0x00, 0x00, 0x00),
        }
        SIZES = {1: 16, 2: 14, 3: 12, 4: 11}
        h = self.doc.add_heading(text, level=min(level, 4))
        if h.runs:
            h.runs[0].font.name = 'Calibri'
            h.runs[0].font.size = Pt(SIZES.get(level, 11))
            h.runs[0].font.color.rgb = COLORS.get(level, RGBColor(0,0,0))

    def para(self, text, space_after=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        _add_mixed_run(p, text)
        return p

    def bullet(self, text, level=0):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(level * 0.5)
        _add_mixed_run(p, text)

    def divider(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('─' * 80)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(7)

    def table(self, headers, rows):
        if not headers:
            return
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Nagłówek
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shd(cell, '2E74B5')

        # Wiersze danych
        for ri, row in enumerate(rows):
            tr = table.rows[ri + 1]
            for ci, val in enumerate(row):
                cell = tr.cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                col = status_color(str(val))
                if col:
                    run.font.color.rgb = col
                if ri % 2 == 1:
                    _shd(cell, 'EAF1FB')

        self.doc.add_paragraph()

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        print(f"Zapisano: {path}")


# ──────────────────────────────────────────────
# PARSER — przetwarza linie MD
# ──────────────────────────────────────────────

def parse_table_row(line):
    """Zwraca listę komórek z wiersza tabeli MD."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def is_separator_row(cells):
    """Czy wiersz jest separatorem nagłówka tabeli (np. |---|---|)."""
    return all(re.match(r'^[-: ]+$', c) for c in cells)


def strip_inline(text):
    """Usuwa znaki formatowania MD z tekstu (dla nagłówków itp.)."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    return text.strip()


def md_to_docx(md_path, out_path):
    b = DocBuilder()
    lines = Path(md_path).read_text(encoding='utf-8').splitlines()

    # Stan parsera
    in_table = False
    table_headers = []
    table_rows = []
    pending_blank = False

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if in_table and table_headers:
            b.table(table_headers, table_rows)
        in_table = False
        table_headers = []
        table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # ── Tabela ──────────────────────────────
        if stripped.startswith('|'):
            cells = parse_table_row(stripped)
            if is_separator_row(cells):
                i += 1
                continue
            if not in_table:
                # Nowa tabela — ta linia jest nagłówkiem
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Pozioma linia (---) ──────────────────
        if re.match(r'^-{3,}$', stripped):
            b.divider()
            i += 1
            continue

        # ── Nagłówki ────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = strip_inline(m.group(2))
            b.heading(text, level)
            i += 1
            continue

        # ── Punktory ────────────────────────────
        m = re.match(r'^(\s*)-\s+(.*)', stripped)
        if m:
            indent = len(m.group(1)) // 2
            b.bullet(m.group(2), level=indent)
            i += 1
            continue

        # ── Numerowane listy ────────────────────
        m = re.match(r'^\d+\.\s+(.*)', stripped)
        if m:
            b.bullet(m.group(1))
            i += 1
            continue

        # ── Blockquote (>) ───────────────────────
        m = re.match(r'^>\s*(.*)', stripped)
        if m:
            p = b.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(strip_inline(m.group(1)))
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue

        # ── Pusta linia ─────────────────────────
        if stripped == '':
            pending_blank = True
            i += 1
            continue

        # ── Zwykły paragraf ─────────────────────
        if pending_blank:
            b.doc.add_paragraph()  # odstęp między akapitami
            pending_blank = False

        b.para(stripped)
        i += 1

    # Flush jeśli tabela na końcu pliku
    flush_table()

    b.save(out_path)


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────

def find_audit_md(directory):
    """Szuka pliku *_audyt.md w podanym katalogu."""
    matches = list(Path(directory).glob('*_audyt.md'))
    if len(matches) == 1:
        return str(matches[0])
    if len(matches) > 1:
        print(f"Znaleziono wiele plików audytu: {[m.name for m in matches]}")
        print("Podaj ścieżkę jako argument.")
        sys.exit(1)
    return None


if __name__ == '__main__':
    if len(sys.argv) >= 2:
        md_path = sys.argv[1]
    else:
        md_path = find_audit_md('.')
        if not md_path:
            print("Użycie: python generate_docx.py <ścieżka_do_audytu.md>")
            sys.exit(1)

    md_path = Path(md_path).resolve()
    if not md_path.exists():
        print(f"Plik nie istnieje: {md_path}")
        sys.exit(1)

    out_path = md_path.with_suffix('.docx')
    print(f"Konwertuję: {md_path.name} → {out_path.name}")
    md_to_docx(str(md_path), str(out_path))
