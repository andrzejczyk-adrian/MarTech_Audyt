# -*- coding: utf-8 -*-
"""
Generuje pliki DOCX z audytu patrizia-aryton
Pliki wejsciowe: 20260407_patrizia-aryton_audyt.md, 20260407_patrizia-aryton_podsumowanie.md
Pliki wyjsciowe: 2x DOCX
"""

import sys, io, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"C:\Users\adria\Documents\AI\AUDYT\MarTech\wykonane\patrizia-aryton"

FILES = {
    "audyt": {
        "input": "20260407_patrizia-aryton_audyt.md",
        "output": "20260407_patrizia-aryton_audyt.docx",
        "title": "Audyt MarTech — Patrizia by Aryton",
        "subtitle": "Raport Szczegolowy | Kwiecien 2026"
    },
    "podsumowanie": {
        "input": "20260407_patrizia-aryton_podsumowanie.md",
        "output": "20260407_patrizia-aryton_podsumowanie.docx",
        "title": "Audyt MarTech — Patrizia by Aryton",
        "subtitle": "Podsumowanie dla Wlasciciela | Kwiecien 2026"
    }
}

def clean_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text)
    return text.strip()

def set_heading_style(paragraph, level=1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    sizes = {1: 16, 2: 13, 3: 11, 4: 10}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    colors = {
        1: RGBColor(0x1a, 0x1a, 0x2e),
        2: RGBColor(0x16, 0x21, 0x3e),
        3: RGBColor(0x0f, 0x3d, 0x66),
        4: RGBColor(0x20, 0x60, 0x80),
    }
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    paragraph.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    paragraph.paragraph_format.space_after = Pt(4)

def add_table_from_md(doc, lines, start_idx):
    table_lines = []
    i = start_idx
    while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip() == ''):
        if lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
        i += 1
    if len(table_lines) < 2:
        return start_idx + 1
    rows = []
    for tl in table_lines:
        cells = [c.strip() for c in tl.strip().strip('|').split('|')]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r'^[-: ]+$', c) for c in r)]
    if not rows:
        return i
    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append('')
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            cell.text = clean_md(cell_text)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.runs[0] if p.runs else p.add_run(cell.text)
            run.font.size = Pt(8.5)
            if ri == 0:
                run.font.bold = True
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1a1a2e')
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    doc.add_paragraph()
    return i

def add_bold_inline(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(clean_md(part))

def md_to_docx(md_path, docx_path, title, subtitle):
    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = s.add_run(subtitle)
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == '':
            i += 1
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            p.add_run(clean_md(line[2:]))
            set_heading_style(p, 1)
            i += 1
            continue
        if line.startswith('## '):
            p = doc.add_paragraph()
            p.add_run(clean_md(line[3:]))
            set_heading_style(p, 2)
            i += 1
            continue
        if line.startswith('### '):
            p = doc.add_paragraph()
            p.add_run(clean_md(line[4:]))
            set_heading_style(p, 3)
            i += 1
            continue
        if line.startswith('#### '):
            p = doc.add_paragraph()
            p.add_run(clean_md(line[5:]))
            set_heading_style(p, 4)
            i += 1
            continue

        if line.strip().startswith('---'):
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(clean_md(line.lstrip('> ')))
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        if line.strip().startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            add_bold_inline(p, line[2:])
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            add_bold_inline(p, re.sub(r'^\d+\. ', '', line))
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Code block — skip or render as monospace
        if line.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(lines[i])
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i += 1
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_bold_inline(p, line)
        for run in p.runs:
            run.font.size = Pt(10)
        i += 1

    doc.save(docx_path)
    print(f"Zapisano: {docx_path}")

# Generate both files
for key, cfg in FILES.items():
    input_path = os.path.join(BASE_DIR, cfg["input"])
    output_path = os.path.join(BASE_DIR, cfg["output"])
    print(f"Generuje: {cfg['output']}...")
    md_to_docx(input_path, output_path, cfg["title"], cfg["subtitle"])

print("\nGotowe! Wygenerowane pliki:")
for key, cfg in FILES.items():
    p = os.path.join(BASE_DIR, cfg["output"])
    if os.path.exists(p):
        size = os.path.getsize(p) // 1024
        print(f"  {cfg['output']} ({size} KB)")
