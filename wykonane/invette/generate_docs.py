# -*- coding: utf-8 -*-
"""
Generuje pliki DOCX i PDF z audytow Invette Google Ads.
Pliki wejsciowe: invette_audyt_zbiorczy.md, invette_audyt_podsumowanie.md
Pliki wyjsciowe: 4 pliki (2x DOCX, 2x PDF)
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r"C:\Users\adria\Documents\AI\AUDYT\MarTech\invette"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_heading_style(paragraph, level=1):
    """Apply heading formatting."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    sizes = {1: 16, 2: 13, 3: 11, 4: 10}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    colors = {
        1: RGBColor(0x1a, 0x1a, 0x2e),
        2: RGBColor(0x16, 0x21, 0x3e),
        3: RGBColor(0x0f, 0x3d, 0x66),
        4: RGBColor(0x20, 0x60, 0x80),
    }
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    paragraph.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    paragraph.paragraph_format.space_after = Pt(4)


def add_table_from_md(doc, lines, start_idx):
    """Parse markdown table starting at lines[start_idx] and add to doc. Returns end index."""
    table_lines = []
    i = start_idx
    while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip() == ''):
        if lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 2:
        return start_idx + 1

    # parse rows
    rows = []
    for tl in table_lines:
        cells = [c.strip() for c in tl.strip().strip('|').split('|')]
        rows.append(cells)

    # remove separator row (---|---...)
    rows = [r for r in rows if not all(re.match(r'^[-: ]+$', c) for c in r)]

    if not rows:
        return i

    ncols = max(len(r) for r in rows)
    # pad short rows
    for r in rows:
        while len(r) < ncols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            cell.text = clean_md(cell_text)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.runs[0] if p.runs else p.add_run(cell.text)
            run.font.size = Pt(8.5)
            if ri == 0:
                run.font.bold = True
                # header background
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1a1a2e')
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    doc.add_paragraph()
    return i


def clean_md(text):
    """Remove markdown formatting."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text.strip()


def set_document_margins(doc):
    """Set document margins to 2cm."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_title_page(doc, title, subtitle, date_str):
    """Add a styled title block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(date_str)
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p3.paragraph_format.space_after = Pt(20)

    # horizontal rule
    doc.add_paragraph('─' * 80)
    doc.add_paragraph()


def md_to_docx(md_path, docx_path, doc_title, doc_subtitle):
    """Convert markdown file to DOCX."""
    doc = Document()
    set_document_margins(doc)

    # default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_title_page(doc, doc_title, doc_subtitle, "Marzec–Kwiecień 2026 | MCC 934-203-1404")

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # skip empty lines → small space
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # headings
        if line.startswith('#### '):
            p = doc.add_paragraph(clean_md(line[5:]))
            set_heading_style(p, 4)
            i += 1
            continue
        if line.startswith('### '):
            p = doc.add_paragraph(clean_md(line[4:]))
            set_heading_style(p, 3)
            i += 1
            continue
        if line.startswith('## '):
            p = doc.add_paragraph(clean_md(line[3:]))
            set_heading_style(p, 2)
            i += 1
            continue
        if line.startswith('# '):
            # skip first H1 (already in title page)
            i += 1
            continue

        # horizontal rule
        if line.strip().startswith('---'):
            doc.add_paragraph('─' * 80)
            i += 1
            continue

        # table
        if line.strip().startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        # unordered list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean_md(line.strip()[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.runs[0].font.size = Pt(9.5)
            i += 1
            continue

        # ordered list
        m = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        if m:
            text = clean_md(m.group(2))
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.runs[0].font.size = Pt(9.5)
            i += 1
            continue

        # blockquote
        if line.strip().startswith('>'):
            text = clean_md(line.strip()[1:].strip())
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(1)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # checkbox
        if line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
            text = clean_md(re.sub(r'- \[.\] ', '', line.strip()))
            checked = '[x]' in line
            p = doc.add_paragraph(('☑ ' if checked else '☐ ') + text)
            p.paragraph_format.left_indent = Cm(0.5)
            p.runs[0].font.size = Pt(9.5)
            i += 1
            continue

        # code block
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # skip closing ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.paragraph_format.left_indent = Cm(1)
            run = p.runs[0] if p.runs else p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x20, 0x60, 0x20)
            continue

        # bold line (standalone **)
        if line.strip().startswith('**') and line.strip().endswith('**') and len(line.strip()) > 4:
            text = clean_md(line.strip())
            p = doc.add_paragraph(text)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(10)
            i += 1
            continue

        # normal paragraph
        text = clean_md(line)
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.size = Pt(10)
        i += 1

    doc.save(docx_path)
    print(f"[OK] DOCX zapisany: {docx_path}")


def docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using LibreOffice or win32com."""
    # Try win32com (MS Word)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
        print(f"✅ PDF zapisany (Word): {pdf_path}")
        return True
    except Exception as e:
        print(f"⚠️ win32com nieudane: {e}")

    # Try LibreOffice
    try:
        import subprocess
        result = subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path), docx_path
        ], capture_output=True, text=True, timeout=60)
        if result.returncode == 0:
            # LibreOffice saves with .pdf extension replacing .docx
            soffice_out = docx_path.replace('.docx', '.pdf')
            if os.path.exists(soffice_out) and soffice_out != pdf_path:
                os.rename(soffice_out, pdf_path)
            print(f"✅ PDF zapisany (LibreOffice): {pdf_path}")
            return True
        else:
            print(f"⚠️ LibreOffice błąd: {result.stderr}")
    except Exception as e:
        print(f"⚠️ LibreOffice nieudane: {e}")

    # Fallback: reportlab simple text
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm

        story = []
        styles = getSampleStyleSheet()

        # Read docx text
        from docx import Document as DocxDoc
        d = DocxDoc(docx_path)
        for para in d.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text[:500], styles['Normal']))
                story.append(Spacer(1, 0.2*cm))

        doc_pdf = SimpleDocTemplate(pdf_path, pagesize=A4,
                                     leftMargin=2.5*cm, rightMargin=2.5*cm,
                                     topMargin=2*cm, bottomMargin=2*cm)
        doc_pdf.build(story)
        print(f"✅ PDF zapisany (reportlab fallback): {pdf_path}")
        return True
    except Exception as e:
        print(f"❌ Reportlab fallback nieudane: {e}")

    return False


# ── main ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    files = [
        {
            'md': os.path.join(OUTPUT_DIR, '20260403_invette_martech_audyt.md'),
            'docx': os.path.join(OUTPUT_DIR, '20260403_invette_martech_audyt.docx'),
            'pdf': os.path.join(OUTPUT_DIR, '20260403_invette_martech_audyt.pdf'),
            'title': 'AUDYT MARTECH — INVETTE MCC',
            'subtitle': 'Raport Zbiorczy | Kwiecień 2026 | 94 konta aktywne',
        },
        {
            'md': os.path.join(OUTPUT_DIR, '20260403_invette_martech_podsumowanie.md'),
            'docx': os.path.join(OUTPUT_DIR, '20260403_invette_martech_podsumowanie.docx'),
            'pdf': os.path.join(OUTPUT_DIR, '20260403_invette_martech_podsumowanie.pdf'),
            'title': 'INVETTE — PODSUMOWANIE AUDYTU MARTECH',
            'subtitle': 'Executive Summary | MCC 934-203-1404 | Kwiecień 2026',
        },
    ]

    for f in files:
        print(f"\n{'='*60}")
        print(f"Przetwarzam: {os.path.basename(f['md'])}")
        md_to_docx(f['md'], f['docx'], f['title'], f['subtitle'])
        docx_to_pdf(f['docx'], f['pdf'])

    print("\n✅ Gotowe! Pliki w:", OUTPUT_DIR)
    print("   20260403_invette_martech_audyt.docx")
    print("   20260403_invette_martech_audyt.pdf")
    print("   20260403_invette_martech_podsumowanie.docx")
    print("   20260403_invette_martech_podsumowanie.pdf")
